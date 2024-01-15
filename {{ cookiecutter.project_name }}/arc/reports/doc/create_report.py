# -*- coding: utf-8 -*-
"""
Module for generating reports in words document and pdf.
"""
import datetime
import json
import logging
import os
import time
from copy import deepcopy
from io import BytesIO

import docx  # noqa
import requests

from docx import table, Document  # noqa
from docx.enum.dml import MSO_THEME_COLOR_INDEX  # noqa
from docx.opc.oxml import parse_xml  # noqa
from docx.oxml import ns, OxmlElement  # noqa
from docx.oxml.ns import nsdecls  # noqa
from docx.shared import RGBColor, Inches, Pt, Cm  # noqa
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # noqa
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # noqa
from arc.contrib.utilities import load_translation
from arc.contrib.tools.formatters import replace_chars
from arc.core.test_method.exceptions import TalosReportException
from arc.reports.html.utils import parse_url_params
from arc.settings.settings_manager import Settings
from arc.core.behave.template_var import replace_template_var

logger = logging.getLogger(__name__)

BASE_PATH = os.path.abspath(os.path.join(os.path.abspath(__file__), os.pardir)) + os.sep
DOC_PATH = os.path.join(Settings.REPORTS_PATH.get(force=True), 'doc') + os.sep

gnu_translations = load_translation('docs_reports')
_ = gnu_translations.gettext
XML_SPACE = 'xml:space'
FLD_CHAR = 'w:fldChar'
FLD_CHAR_TYPE = 'w:fldCharType'


class CreateDOC:
    """
    Report generation class in docx format.
    """
    document: Document
    scenario_name = ""
    scenario_location = ""
    application_name = ""
    drive_type = ''
    summary_table: table
    total_step_table: table
    cell_color = r'<w:shd {} w:fill="f6f6f6"/>'
    reports = []

    def generate_document_report(self, feature, global_data):
        """
        Generate document report depending on the global data.
        """
        logger.debug(f'Generating document docx')
        load_translation('docs_reports')
        logger.debug('Babel translation for documents reports loaded')
        self.drive_type = feature['driver']
        self.application_name = global_data['application']

        for scenario in feature['elements']:
            if scenario["type"] == 'scenario':
                logger.debug(f"Generating document docx evidence for scenario: {scenario['name']}")
                self.document = Document()
                self.scenario_name = scenario['name']
                self.scenario_location = scenario['location']
                self.__header()
                self.__add_page_number(self.document.sections[0].footer.paragraphs[0])
                self.document.add_paragraph('')
                testcase_name = self.scenario_name + ' - ' + str(self.drive_type).upper()
                self.__set_scenario_feature_name(f"{_('Feature')}: {feature['name']}")
                feature_description = ' '.join(feature.get('description', ''))
                self.__set_scenario_feature_name(f"{_('Scenario')}: {testcase_name}")
                scenario_description = ' '.join(scenario.get('description', ''))
                self.__generate_global_data_table(global_data)
                self.__generate_summary_table(scenario)
                self.__generate_description_table(f"{_('Feature description')}:", feature_description)
                self.__generate_description_table(f"{_('Scenario description')}:", scenario_description)
                self.__generate_total_step_table(scenario)
                self.__generate_doc_body(scenario)
                self.__end()

        return self.reports

    def __set_scenario_feature_name(self, text):
        """
        Set in document scenario feature text name.
        """
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run(text)
        font = run.font
        font.color.rgb = RGBColor(0, 0, 0)
        font.bold = True
        font.size = Pt(12)

    def __header(self):
        """
        Set in document header
        """
        section = self.document.sections[0]
        header = section.header
        table_header = header.add_table(1, 3, Inches(6))
        table_header.style = 'TableGrid'
        cell = table_header.rows[0].cells[0]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture('arc/web/static/images/taloslogo.png', width=Inches(0.4))
        cell = table_header.cell(0, 1)
        cell.text = self.application_name
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell = table_header.cell(0, 2)
        cell.text = datetime.datetime.today().strftime('%d/%m/%Y')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell in table_header.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in table_header.columns[1].cells:
            cell.width = Inches(4.5)
        for cell in table_header.columns[2].cells:
            cell.width = Inches(1.0)

    @staticmethod
    def __create_element(name):
        """
        Create a document element
        """
        return OxmlElement(name)

    @staticmethod
    def __create_attribute(element, name, value):
        """
        Create a element attribute in document with name and value.
        """
        element.set(ns.qn(name), value)

    def __add_page_number(self, paragraph):
        """
        Add page number.
        """
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        page_run = paragraph.add_run()
        t1 = self.__create_element('w:t')
        self.__create_attribute(t1, XML_SPACE, 'preserve')  # noqa
        t1.text = 'Page '
        page_run._r.append(t1)  # noqa

        page_num_run = paragraph.add_run()

        field_char_1 = self.__create_element(FLD_CHAR)  # noqa
        self.__create_attribute(field_char_1, FLD_CHAR_TYPE, 'begin')  # noqa

        instr_text = self.__create_element('w:instrText')  # noqa
        self.__create_attribute(instr_text, XML_SPACE, 'preserve')
        instr_text.text = "PAGE"

        fld_char2 = self.__create_element(FLD_CHAR)
        self.__create_attribute(fld_char2, FLD_CHAR_TYPE, 'end')

        page_num_run._r.append(field_char_1)  # noqa
        page_num_run._r.append(instr_text)  # noqa
        page_num_run._r.append(fld_char2)  # noqa

        of_run = paragraph.add_run()
        t2 = self.__create_element('w:t')
        self.__create_attribute(t2, XML_SPACE, 'preserve')
        t2.text = ' of '
        of_run._r.append(t2)  # noqa

        fld_char3 = self.__create_element(FLD_CHAR)
        self.__create_attribute(fld_char3, FLD_CHAR_TYPE, 'begin')

        instr_text2 = self.__create_element('w:instrText')
        self.__create_attribute(instr_text2, XML_SPACE, 'preserve')
        instr_text2.text = "NUMPAGES"

        fld_char4 = self.__create_element(FLD_CHAR)
        self.__create_attribute(fld_char4, FLD_CHAR_TYPE, 'end')

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fld_char3)  # noqa
        num_pages_run._r.append(instr_text2)  # noqa
        num_pages_run._r.append(fld_char4)  # noqa

    def __generate_global_data_table(self, global_data):
        """
        Generate global data table in document.
        """
        table_r = self.document.add_table(rows=3, cols=2)
        table_r.style = 'TableGrid'
        cell = table_r.cell(0, 0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table_r.cell(0, 0).text = _('Executed by user')
        table_r.cell(1, 0).text = _('Business area')
        table_r.cell(2, 0).text = _('Entity')

        cell = table_r.cell(0, 1)
        cell.text = str(global_data['user_code']).capitalize()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(1, 1)
        cell.text = global_data['business_area']
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(2, 1)
        cell.text = str(global_data['entity'])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table_r.cell(1, 0)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(1, 1)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa

        self.document.add_paragraph('')

    def __generate_summary_table(self, scenario):
        """
        Generate summary table in document
        """
        table_r = self.document.add_table(rows=11, cols=2)
        self.summary_table = table_r
        table_r.style = 'TableGrid'
        cella = table_r.cell(0, 0)
        cellb = table_r.cell(0, 1)
        cella.merge(cellb)
        cell = table_r.cell(0, 0)
        self.add_bookmark(cell.paragraphs[0], _('Test Case Execution Summary'), "ResumenTable")
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.__set_result_format(table_r)

        table_r.cell(1, 0).text = _('Result')
        table_r.cell(2, 0).text = _('Number of steps')
        table_r.cell(3, 0).text = _('Correct steps')
        table_r.cell(4, 0).text = _('Incorrect steps')
        table_r.cell(5, 0).text = _('No Run steps')
        table_r.cell(6, 0).text = _('Start')
        table_r.cell(7, 0).text = _('End')
        table_r.cell(8, 0).text = _('Duration')
        table_r.cell(9, 0).text = _('Execution Type')
        table_r.cell(10, 0).text = _('Environment')

        if str(scenario['status']) == 'passed':
            result = 'Passed'
        elif str(scenario['status']) == "failed":
            result = 'Failed'
        else:
            result = 'No Run'

        steps = []
        steps_passed = []
        steps_failed = []
        steps_skip = []
        for step in scenario['steps']:
            steps.append(step)
            if step.get('result'):
                if step['result']['status'] == 'passed':
                    steps_passed.append(step)
                elif step['result']['status'] == 'failed':
                    steps_failed.append(step)
            else:
                steps_skip.append(step)

        no_run = steps.__len__() - steps_passed.__len__() - steps_failed.__len__()
        cell = table_r.cell(1, 1)
        cell._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        cell.text = _(result)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(2, 1)
        cell.text = steps.__len__().__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(3, 1)
        cell.text = steps_passed.__len__().__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(4, 1)
        cell.text = steps_failed.__len__().__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(5, 1)
        cell.text = no_run.__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(6, 1)
        cell.text = datetime.datetime.fromtimestamp(scenario.get('start_time', 0)).strftime("%d/%m/%Y %H:%M:%S")
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(7, 1)
        cell.text = datetime.datetime.fromtimestamp(scenario.get('end_time', 0)).strftime("%d/%m/%Y %H:%M:%S")
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(8, 1)
        test_duration = datetime.timedelta(seconds=scenario.get('duration', 0))
        cell.text = test_duration.__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(9, 1)
        cell.text = str(self.drive_type).capitalize()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(10, 1)
        cell.text = str(Settings.PYTALOS_PROFILES.get('environment')).upper()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table_r.cell(1, 0)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(1, 1)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(3, 0)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(3, 1)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(5, 0)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(5, 1)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(7, 0)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(7, 1)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(9, 0)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        table_r.cell(9, 1)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        self.document.add_paragraph('')
        self.__set_result_format(self.summary_table)

    def __generate_description_table(self, title, description):
        """
        Generate a description in the table.
        """
        table_r = self.document.add_table(rows=2, cols=1)
        self.summary_table = table_r
        table_r.style = 'TableGrid'
        cell = table_r.cell(0, 0)
        self.add_bookmark(cell.paragraphs[0], title, "ResumenTable")
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.__set_result_format(table_r)
        table_r.cell(1, 0).text = description
        table_r.cell(1, 0)._tc.get_or_add_tcPr().append(parse_xml(self.cell_color.format(nsdecls('w'))))  # noqa
        self.document.add_paragraph('')
        self.__set_result_format(self.summary_table)

    def __generate_total_step_table(self, scenario):
        """
        Generate total data information step table.
        """
        steps = []
        for step in scenario['steps']:
            if not step.get('result'):
                step['result'] = {"status": 'skipped', 'duration': '-'}
            steps.append(step)

        rows = steps.__len__()
        table_r = self.document.add_table(rows=rows + 1, cols=4)
        self.total_step_table = table_r
        table_r.style = 'TableGrid'

        cell = table_r.cell(0, 0)
        cell.text = _('Steps')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cell = table_r.cell(0, 1)
        cell.text = _('Description')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cell = table_r.cell(0, 2)
        cell.text = _('Result')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cell = table_r.cell(0, 3)
        cell.text = _('Duration')
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cont_row = 1
        for row in steps:
            cell = table_r.cell(cont_row, 0)
            step_text = f"{_('Step')} " + str(cont_row)
            self.add_link(cell.paragraphs[0], f"{_('Step')}" + str(cont_row) + "bm", step_text)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell = table_r.cell(cont_row, 1)
            cell.text = row['name']
            cell = table_r.cell(cont_row, 2)
            cell.text = self.__format_status(row)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell = table_r.cell(cont_row, 3)
            if isinstance(row['result']['duration'], float):
                test_duration = time.strftime('%H:%M:%S', time.gmtime(row['result']['duration']))
            else:
                test_duration = row['result']['duration']
            cell.text = test_duration.__str__()
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            if cont_row % 2 != 0:
                table_r.cell(cont_row, 0)._tc.get_or_add_tcPr().append(  # noqa
                    parse_xml(self.cell_color.format(nsdecls('w')))
                )
                table_r.cell(cont_row, 1)._tc.get_or_add_tcPr().append(  # noqa
                    parse_xml(self.cell_color.format(nsdecls('w')))
                )
                table_r.cell(cont_row, 2)._tc.get_or_add_tcPr().append(  # noqa
                    parse_xml(self.cell_color.format(nsdecls('w')))
                )
                table_r.cell(cont_row, 3)._tc.get_or_add_tcPr().append(  # noqa
                    parse_xml(self.cell_color.format(nsdecls('w')))
                )
            cont_row += 1
        self.__set_column_width(table_r)
        self.__set_result_format(table_r)

    @staticmethod
    def __set_column_width(table_width):
        """
        Set column width from table width.
        """
        for cell in table_width.columns[0].cells:
            cell.width = Inches(1.0)
        for cell in table_width.columns[1].cells:
            cell.width = Inches(3.0)
        for cell in table_width.columns[2].cells:
            cell.width = Inches(1.0)
        for cell in table_width.columns[3].cells:
            cell.width = Inches(1.0)

    @staticmethod
    def __set_result_format(table_in):
        """
        Format result and text in a table passed by parameter.
        """
        for row in table_in.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text == _('Failed'):
                            font = run.font
                            font.color.rgb = RGBColor(252, 3, 3)
                            font.bold = True
                        if run.text == _('Passed'):
                            font = run.font
                            font.color.rgb = RGBColor(3, 156, 51)
                            font.bold = True
                        if run.text == _('Skipped'):
                            font = run.font
                            font.color.rgb = RGBColor(0, 108, 255)
                            font.bold = True
                        if run.text == _('Error'):
                            font = run.font
                            font.color.rgb = RGBColor(191, 13, 13)
                            font.bold = True
                        if run.text == _('Description:') or run.text == _('Expected result:') or \
                                run.text == _('Obtained result:') or run.text == _('Data Table'):
                            font = run.font
                            font.color.rgb = RGBColor(0, 0, 0)
                            font.bold = True
                        if _("Nota: ") in run.text:
                            font = run.font
                            font.color.rgb = RGBColor(0, 0, 0)
                            font.bold = True
                        else:
                            font = run.font
                            font.bold = False

    def __end(self):
        """
        Finish creating the document
        """
        path = ''
        try:
            # Adjust table width
            find = False
            for current_table in self.document.tables:
                if len(current_table.rows) == 1:
                    find = True
                if find is False:
                    current_table.autofit = False
                    current_table.allow_autofit = False
                find = False
            scenario_name = '%.100s' % self.scenario_name
            path = self.__get_doc_path(replace_chars(scenario_name))
            if os.path.exists(path):
                basename, ext = os.path.splitext(path)
                index_location = str(self.scenario_location).rfind(':')
                path = basename + f"_{self.scenario_location[index_location+1:]}" + ext
            self.document.save(path)
            self.reports.append(path)
            logger.debug(f'Document docx generated in path: {path}')
        except(PermissionError,):
            raise TalosReportException(
                'There was an error creating the evidence word document:\n'
                f'document word path: {path}\n'
                'Maybe you have the document open?'
            )

    def __get_doc_path(self, scenario_name):
        """
        Get document path from scenario name.
        """
        path = str(DOC_PATH) + str(self.drive_type).upper() + "-" + scenario_name + ".docx"
        return path

    @staticmethod
    def __format_status(step):
        """
        Format step result in document.
        """
        if step.get('result'):
            if str(type(step)).__contains__('dict'):
                if step['result']['status'] == 'passed':
                    return _('Passed')
                if step['result']['status'] == 'failed':
                    return _('Failed')
                else:
                    return _('Skipped')
            else:
                if step.status.__str__() == 'Status.passed':
                    return _('Passed')
                if step.status.__str__() == 'Status.failed':
                    return _('Failed')
                else:
                    return _('Skipped')

    def __generate_doc_body(self, scenario):
        """
        Generate document step information from scenario.
        """
        cont = 1
        for step in scenario['steps']:
            if step['result']['status'] != 'skipped':
                self.document.add_page_break()
                self.document.add_paragraph('')
                table_r = self.document.add_table(rows=1, cols=4)
                table_r.style = 'TableGrid'
                cell = table_r.cell(0, 0)
                step_text = "Step " + cont.__str__()
                self.add_bookmark(cell.paragraphs[0], step_text, _("Step") + cont.__str__() + "bm")
                cont += 1
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                cell = table_r.cell(0, 1)
                cell.text = _(str(step.get('result').get('status')).capitalize())
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                cell = table_r.cell(0, 2)
                cell.text = datetime.datetime.fromtimestamp(step.get("start_time")).strftime("%d/%m/%Y %H:%M:%S")
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                cell = table_r.cell(0, 3)
                self.add_link(cell.paragraphs[0], "ResumenTable", _("Go to Summary Table"))
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for cell in table_r.columns[0].cells:
                    cell.width = Inches(1.0)
                for cell in table_r.columns[1].cells:
                    cell.width = Inches(1.0)
                for cell in table_r.columns[2].cells:
                    cell.width = Inches(3.0)
                for cell in table_r.columns[3].cells:
                    cell.width = Inches(3.0)

                self.__set_result_format(table_r)
                self.__style_headers(f"{_('Description')}:")
                self.document.add_paragraph(f'''{step['step_type']} {step['name']}''')

                if step.get('text'):
                    paragraph = self.document.add_paragraph()
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    run = paragraph.add_run(step.get('text'))
                    font = run.font
                    font.italic = True

                # Add table
                if step.get('table'):
                    table_t = self.document.add_table(rows=len(step['table']['rows']) + 1,
                                                      cols=len(step['table']['headings']))
                    table_t.style = 'TableGrid'
                    count_header = 0
                    for heading in step['table']['headings']:
                        cell = table_t.cell(0, count_header)
                        cell.text = replace_template_var(heading)
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        count_header += 1

                    count_row = 1
                    for rows in step['table']['rows']:
                        for j, value in enumerate(rows):
                            cell = table_t.cell(count_row, j)
                            cell.text = replace_template_var(value)
                            if count_row % 2 != 0:
                                table_t.cell(count_row, j)._tc.get_or_add_tcPr().append(  # noqa
                                    parse_xml(self.cell_color.format(nsdecls('w')))
                                )
                        count_row += 1
                self.__style_headers(f"{_('Expected result')}:")
                self.document.add_paragraph(step.get('result').get('expected_result'))
                self.__style_headers(f"{_('Obtained result')}:")
                self.document.add_paragraph(step.get('result').get('obtained_result'))
                if step.get('additional_text'):
                    self.__style_headers(f"{_('Additional text')}:")
                    self.document.add_paragraph('\n'.join(step.get('additional_text')))
                self.document.add_paragraph('')
                self.__add_api_evidence_body(step)
                try:
                    for screenshot in step['screenshots']:
                        if str(self.drive_type).lower() == "android" or str(self.drive_type).lower() == "ios":
                            self.document.add_picture(screenshot, width=Cm(10))
                            last_paragraph = self.document.paragraphs[-1]
                            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            self.document.add_picture(screenshot, width=Cm(18))

                except FileNotFoundError:
                    self.document.add_paragraph('')
                self.document.add_paragraph('')

    def __style_headers(self, text):
        """
        Format headers with style.
        """
        self.document.add_paragraph('')
        run = self.document.add_paragraph().add_run(text)
        font = run.font
        font.color.rgb = RGBColor(0, 0, 0)
        font.bold = True
        font.italic = True
        font.size = Pt(14)

    def __add_api_evidence_body(self, step):
        """
        Add api evidences body to document from step.
        """
        if step.get('jsons'):
            for inte in step.get("jsons"):
                if inte and str(inte) != "null":
                    api_evidence = inte
                    table_r = self.document.add_table(rows=2, cols=1)
                    table_r.style = 'TableGrid'
                    cell = table_r.cell(0, 0)
                    cell.text = api_evidence['title']

                    cell = table_r.cell(1, 0)
                    table_r.cell(1, 0)._tc.get_or_add_tcPr().append(  # noqa
                        parse_xml(self.cell_color.format(nsdecls('w')))
                    )
                    cell.text = json.dumps(api_evidence['content'], indent=4)
                    self.set_api_evidence_title_colour(table_r)
                    self.document.add_paragraph('')

        if step.get('api_info'):
            table_r = self.document.add_table(rows=7, cols=2)
            table_r.style = 'TableGrid'
            cella = table_r.cell(0, 0)
            cellb = table_r.cell(0, 1)
            cell_merge = cella.merge(cellb)
            cell_merge.text = 'Request Info'

            info_dict = step.get("api_info")
            cell = table_r.cell(1, 0)
            cell.text = "URL"
            cell = table_r.cell(2, 0)
            cell.text = _("Method")
            cell = table_r.cell(3, 0)
            cell.text = _("Status Code")
            cell = table_r.cell(4, 0)
            cell.text = _("Remote Address")
            cell = table_r.cell(5, 0)
            cell.text = _("Request headers")
            cell = table_r.cell(6, 0)
            cell.text = _("Request body")

            cell = table_r.cell(1, 1)
            url = info_dict.get("url")
            if params := info_dict.get('params'):
                url = parse_url_params(url, params)
            cell.text = url
            cell = table_r.cell(2, 1)
            cell.text = str(info_dict.get("method")).upper()
            cell = table_r.cell(3, 1)
            cell.text = str(info_dict.get("status_code")) + " " + info_dict.get("reason")
            cell = table_r.cell(4, 1)
            cell.text = str(info_dict.get("remote Address"))
            cell = table_r.cell(5, 1)

            headers = json.dumps(info_dict.get("headers"), indent=4).replace('\\n', '$n')
            headers = headers.replace('\\', '')
            headers = headers.replace('"', '')
            headers = headers.replace('$n', '\n')
            cell.text = str(headers)
            cell = table_r.cell(6, 1)

            body = json.dumps(info_dict.get("body"), indent=4).replace('\\n', '$n')
            body = body.replace('\\', '')
            body = body.replace('"', '')
            body = body.replace('$n', '\n')
            cell.text = str(body)
            table_r.cell(1, 0)._tc.get_or_add_tcPr().append(  # noqa
                parse_xml(self.cell_color.format(nsdecls('w')))
            )
            table_r.cell(1, 1)._tc.get_or_add_tcPr().append(  # noqa
                parse_xml(self.cell_color.format(nsdecls('w')))
            )
            table_r.cell(3, 0)._tc.get_or_add_tcPr().append(  # noqa
                parse_xml(self.cell_color.format(nsdecls('w')))
            )
            table_r.cell(3, 1)._tc.get_or_add_tcPr().append(  # noqa
                parse_xml(self.cell_color.format(nsdecls('w')))
            )

            for cell in table_r.columns[0].cells:
                cell.width = Inches(1.5)
            for cell in table_r.columns[1].cells:
                cell.width = Inches(5.5)
            self.set_api_evidence_title_colour(table_r)
            self.document.add_paragraph('')

            if step.get("response_headers"):
                table_r = self.document.add_table(rows=2, cols=1)
                table_r.style = 'TableGrid'
                cell = table_r.cell(0, 0)
                cell.text = _("Request headers")

                cell = table_r.cell(1, 0)
                response = json.dumps(step.get('response_headers'), indent=4).replace('\\n', '$n')
                response = response.replace('\\', '')
                response = response.replace('"', '')
                response = response.replace('$n', '\n')
                cell.text = response
                table_r.cell(1, 0)._tc.get_or_add_tcPr().append(  # noqa
                    parse_xml(self.cell_color.format(nsdecls('w')))
                )
                self.set_api_evidence_title_colour(table_r)
                self.document.add_paragraph('')

            if step.get("response_content"):
                table_r = self.document.add_table(rows=2, cols=1)
                table_r.style = 'TableGrid'
                cell = table_r.cell(0, 0)
                cell.text = _("Request Response")

                cell = table_r.cell(1, 0)
                response = json.dumps(step.get('response_content'), indent=4).replace('\\n', '$n')
                response = response.replace('\\', '')
                response = response.replace('"', '')
                response = response.replace('$n', '\n')
                cell.text = response
                table_r.cell(1, 0)._tc.get_or_add_tcPr().append(  # noqa
                    parse_xml(self.cell_color.format(nsdecls('w')))
                )
                self.set_api_evidence_title_colour(table_r)
                self.document.add_paragraph('')

        if step.get('unit_tables'):
            for dictionary in step.get("unit_tables"):
                dictionary = deepcopy(dictionary)
                len_dict = len(dictionary)
                row_cont = len_dict

                table_r = self.document.add_table(rows=row_cont, cols=2)
                table_r.style = 'TableGrid'
                cella = table_r.cell(0, 0)
                cellb = table_r.cell(0, 1)
                cell_merge = cella.merge(cellb)
                cell_merge.text = dictionary["title"]
                del dictionary["title"]
                cont = 1
                for key in dictionary.keys():
                    if key != "error":
                        if dictionary.get(key) is not None:
                            if cont % 2 != 0:
                                table_r.cell(cont, 0)._tc.get_or_add_tcPr().append(  # noqa
                                    parse_xml(self.cell_color.format(nsdecls('w')))
                                )
                                table_r.cell(cont, 1)._tc.get_or_add_tcPr().append(  # noqa
                                    parse_xml(self.cell_color.format(nsdecls('w')))
                                )
                            cell = table_r.cell(cont, 0)
                            cell.text = key
                            cell = table_r.cell(cont, 1)
                            cell.text = str(dictionary.get(key))
                            cont += 1
                    elif key == "error" and dictionary.get(key) is not None:
                        if cont % 2 != 0:
                            table_r.cell(cont, 0)._tc.get_or_add_tcPr().append(  # noqa
                                parse_xml(self.cell_color.format(nsdecls('w')))
                            )
                            table_r.cell(cont, 1)._tc.get_or_add_tcPr().append(  # noqa
                                parse_xml(self.cell_color.format(nsdecls('w')))
                            )
                        cell = table_r.cell(cont, 0)
                        cell.text = str(key)
                        cell = table_r.cell(cont, 1)
                        cell.text = str(dictionary.get(key))
                        cont += 1
                self.__set_result_format(table_r)
                self.set_api_evidence_title_colour(table_r)
                self.document.add_paragraph('')

    @staticmethod
    def set_api_evidence_title_colour(table_in):
        """
        Format with color api evidence title.
        """
        header = True
        for row in table_in.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if header:
                            font = run.font
                            font.color.rgb = RGBColor(252, 3, 3)
                            font.bold = True
                            header = False

    @staticmethod
    def add_bookmark(paragraph, bookmark_text, bookmark_name):
        """
        Add new bookmark from text and name.
        """
        run = paragraph.add_run()
        tag = run._r  # noqa
        start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
        start.set(docx.oxml.ns.qn('w:id'), '0')
        start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(start)

        text = docx.oxml.OxmlElement('w:r')
        text.text = bookmark_text
        tag.append(text)

        end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
        end.set(docx.oxml.ns.qn('w:id'), '0')
        end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(end)

    @staticmethod
    def add_link(paragraph, link_to, text, tool_tip=None):
        """
        Add link from text and target anchor
        """
        # create hyperlink node
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )

        if tool_tip is not None:
            # set attribute for link to bookmark
            hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip, )

        new_run = docx.oxml.shared.OxmlElement('w:r')
        r_pr = docx.oxml.shared.OxmlElement('w:rPr')
        new_run.append(r_pr)
        new_run.text = text
        hyperlink.append(new_run)
        r = paragraph.add_run()
        r._r.append(hyperlink)  # noqa
        r.font.name = "Calibri"
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True
